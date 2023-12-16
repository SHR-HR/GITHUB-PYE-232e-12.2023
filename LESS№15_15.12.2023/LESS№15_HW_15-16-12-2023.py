# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
'''''   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓   ↓
Дата выполнения домашней работы: 15-16 - ДЕКАБРЯ 2023
'''''
'''
Курс: Разработка Web-приложений на Python, с применением Фреймворка Django
Дисциплина: Основы программирования на Python
'''
'''
Урок от 15.12.2023
Домашняя работа №15: Работа с комплексными файлами - excel, json, word. Библиотеки openpyxl, json, docx
'''
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
'''
Выполните следующие задания:
'''
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
'''
Задание №1

а) Создайте word файл в операционной системе, заполните его текстом «Hello Python».
б) Прочитайте этот файл, только жирный текст в python-строку и выведите на экран.
в) Создайте абзац с текстом и запишите это в новый word-файл, изменит шрифт и размер шрифта.
'''
'''
Решение:
'''
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
'''
а) Создайте word файл в операционной системе, заполните его текстом «Hello Python». ↓ 
'''
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
!!! pip install python-docx !!!
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
from docx import Document
from docx.shared import Pt

# Создаем новый документ
doc = Document()

# Добавляем параграф с текстом "Hello Python"
paragraph = doc.add_paragraph("Hello ")
run = paragraph.add_run("Python")
run.bold = True  # Устанавливаем стиль "жирный"

# Сохраняем документ
doc.save("hello_python.docx")
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
'''
1: Импорт библиотеки и классов:
'''
from docx import Document
from docx.shared import Pt
'''
Эта строка кода импортирует класс Document из библиотеки python-docx для создания и редактирования документов Word.
Pt из docx.shared используется для задания размера шрифта.
'''
'''
2: Создание нового документа:
'''
doc = Document()
'''
Создается новый объект Document, представляющий собой пустой документ Word.
'''
'''
3: Добавление параграфа с текстом "Hello Python":
'''
paragraph = doc.add_paragraph("Hello ")
run = paragraph.add_run("Python")
'''
Создается новый параграф, и в этот параграф добавляется текст "Hello ".
Затем создается объект run (запуск), который представляет собой последовательность текста с общими свойствами
форматирования. В данном случае, к тексту "Python" добавляется с использованием add_run.
'''
'''
4: Установка стиля "жирный" для текста "Python":
'''
run.bold = True
'''
Эта строка устанавливает стиль "жирный" для текста, добавленного с использованием объекта run.
Теперь слово "Python" будет отображаться в жирном стиле.
'''
'''
5: Сохранение документа:
'''
doc.save("hello_python.docx")
'''
Сохраняет созданный документ в файл с именем "hello_python.docx".
Если файл с таким именем уже существует, он будет перезаписан. Если нет, то будет создан новый файл.
'''
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
'''
б) Прочитайте этот файл, только жирный текст в python-строку и выведите на экран. ↓ 
'''
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
from docx import Document

# Открываем существующий документ
doc = Document("hello_python.docx")

# Получаем жирный текст из документа
bold_text = ""
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text

# Выводим жирный текст на экран
print("Жирный текст из файла:")
print(bold_text)
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
'''
1: Импорт библиотеки и класса:
'''
from docx import Document
'''
Эта строка кода импортирует класс Document из библиотеки python-docx. Он используется для работы с документами Word.
'''
'''
2: Открытие существующего документа:
'''
doc = Document("hello_python.docx")
'''
Создается объект Document, представляющий существующий документ Word с именем "hello_python.docx".
Это предполагает, что в рабочей директории программы есть файл с таким именем.
'''
'''
3: Извлечение жирного текста из документа:
'''
bold_text = ""
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text
'''
Происходит итерация по всем параграфам в документе и внутри каждого параграфа - по всем "запускам" (runs).
Если у запуска run установлен стиль "жирный" (run.bold == True), то текст этого запуска добавляется
к переменной bold_text.
'''
'''
4: Вывод жирного текста на экран:
'''
print("Жирный текст из файла:")
print(bold_text)
'''
Выводит текст, который был помечен как жирный, на экран.
'''
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
'''
в)  Создайте абзац с текстом, измените шрифт и размер шрифта, и запишите это в новый Word файл:. ↓ 
'''
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Создаем новый документ
new_doc = Document()

# Добавляем абзац с текстом
paragraph = new_doc.add_paragraph("Новый абзац с измененным шрифтом и размером шрифта.")

# Изменяем шрифт и размер шрифта
run = paragraph.runs[0]
run.font.size = Pt(14)
run.font.name = "Arial"  # Устанавливаем шрифт "Arial"

# Изменяем выравнивание абзаца (по центру)
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Сохраняем новый документ
new_doc.save("new_document_formatted_v2.docx")

print("Новый документ с измененным форматированием создан и сохранен.")
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
'''
1: Импорт библиотек и классов:
'''
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
'''
Эта строка кода импортирует класс Document для создания и редактирования документов Word,
Pt для задания размера шрифта, и WD_PARAGRAPH_ALIGNMENT для определения выравнивания абзаца.
'''
'''
2: Создание нового документа:
'''
new_doc = Document()
'''
Создается новый объект Document, представляющий собой пустой документ Word.
'''
'''
3: Добавление абзаца с текстом:
'''
paragraph = new_doc.add_paragraph("Новый абзац с измененным шрифтом и размером шрифта.")
'''
Создается новый абзац и добавляется к документу с текстом "Новый абзац с измененным шрифтом и размером шрифта."
'''
'''
4: Изменение шрифта и размера шрифта:
'''
run = paragraph.runs[0]
run.font.size = Pt(14)
run.font.name = "Arial"  # Устанавливаем шрифт "Arial"
'''
Получается первый запуск (run) в добавленном абзаце и устанавливается размер шрифта (14 пунктов) и шрифт (Arial).
'''
'''
5: Изменение выравнивания абзаца (по центру):
'''
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
'''
Устанавливается выравнивание абзаца по центру с использованием WD_PARAGRAPH_ALIGNMENT.
'''
'''
6: Сохранение нового документа:
'''
new_doc.save("new_document_formatted_v2.docx")
'''
Сохраняет созданный документ в файл с именем "new_document_formatted_v2.docx".
'''
'''
7: Вывод уведомления на экран:
'''
print("Новый документ с измененным форматированием создан и сохранен.")
'''
Выводит сообщение о создании и сохранении нового документа с измененным форматированием.
'''
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
""" ВСЕ СРАЗУ """
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Часть а) Создаем Word файл и заполняем его текстом "Hello Python"
doc = Document()
paragraph = doc.add_paragraph("Hello ")
run = paragraph.add_run("Python")
run.bold = True  # Устанавливаем стиль "жирный"
doc.save("hello_python.docx")

# Часть б) Открываем существующий документ и выводим жирный текст
read_doc = Document("hello_python.docx")
bold_text = ""
for paragraph in read_doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text
print("Жирный текст из файла:")
print(bold_text)

# Часть в) Создаем новый документ с абзацем, изменяем шрифт и размер шрифта
new_doc = Document()
new_paragraph = new_doc.add_paragraph("Новый абзац с измененным шрифтом и размером шрифта.")
new_run = new_paragraph.runs[0]
new_run.font.size = Pt(14)
new_run.font.name = "Arial"  # Устанавливаем шрифт "Arial"
new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
new_doc.save("new_document_formatted.docx")

print("Новый документ с измененным форматированием создан и сохранен.")
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
'''
а) Создаем Word файл и заполняем его текстом "Hello Python"
'''
'''
1: Импорт библиотек и классов:
'''
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
'''
Импортируются необходимые классы и элементы из библиотеки python-docx.
'''
'''
2: Создание нового документа и добавление текста:
'''
doc = Document()
paragraph = doc.add_paragraph("Hello ")
run = paragraph.add_run("Python")
run.bold = True
'''
Создается новый объект Document, к которому добавляется абзац с текстом "Hello " и отдельный "запуск" (run) с
текстом "Python". Устанавливается стиль "жирный" для этого запуска.
'''
'''
3: Сохранение документа:
'''
doc.save("hello_python.docx")
'''
Сохраняет созданный документ в файл "hello_python.docx".
'''
'''
б) Открываем существующий документ и выводим жирный текст:
'''
'''
1: Открытие существующего документа и извлечение жирного текста:
'''
read_doc = Document("hello_python.docx")
bold_text = ""
for paragraph in read_doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text
'''
Открывает существующий документ "hello_python.docx" и извлекает текст, который был отмечен как жирный.
'''
'''
2: Вывод жирного текста на экран:
'''
print("Жирный текст из файла:")
print(bold_text)
'''
Выводит на экран текст, который был помечен как жирный в открытом документе.
'''
'''
в) Создаем новый документ с абзацем, изменяем шрифт и размер шрифта:
'''
'''
1: Создание нового документа и добавление абзаца:
'''
new_doc = Document()
new_paragraph = new_doc.add_paragraph("Новый абзац с измененным шрифтом и размером шрифта.")
'''
Создается новый документ, к которому добавляется абзац с текстом "Новый абзац с измененным шрифтом и размером шрифта."
'''
'''
2: Изменение шрифта и размера шрифта:
'''
new_run = new_paragraph.runs[0]
new_run.font.size = Pt(14)
new_run.font.name = "Arial"
'''
Получается первый запуск в новом абзаце и изменяются его размер и шрифт.
'''
'''
3: Изменение выравнивания абзаца (по центру):
'''
new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
'''
Устанавливается выравнивание абзаца по центру.
'''
'''
4: Сохранение нового документа:
'''
new_doc.save("new_document_formatted.docx")
'''
Сохраняет созданный документ в файл "new_document_formatted.docx".
'''
'''
5: Вывод уведомления на экран:
'''
print("Новый документ с измененным форматированием создан и сохранен.")
'''
Выводит уведомление о создании и сохранении нового документа с измененным форматированием.
'''
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
""" BONUS """
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# Часть а) Создаем Word файл и заполняем его текстом "Hello Python"
doc = Document()
paragraph = doc.add_paragraph("Hello ")
run = paragraph.add_run("Python")
run.bold = True
doc.save("hello_python.docx")

# Часть б) Открываем существующий документ и выводим жирный текст
read_doc = Document("hello_python.docx")
bold_text = ""
for paragraph in read_doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text
print("Жирный текст из файла:")
print(bold_text)

# Часть в) Создаем новый документ с абзацем, изменяем шрифт и размер шрифта
new_doc = Document()
new_paragraph = new_doc.add_paragraph("Новый абзац с измененным шрифтом и размером шрифта.")
new_run = new_paragraph.runs[0]
new_run.font.size = Pt(14)
new_run.font.name = "Arial"
new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
new_doc.save("new_document_formatted.docx")

# Создаем PDF документ
def create_pdf(file_name, text):
    pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))  # Замените 'arial.ttf' на путь к файлу вашего шрифта Arial
    c = canvas.Canvas(file_name, pagesize=letter)
    width, height = letter
    c.setFont("Arial", 14)
    c.drawString(width / 2, height - 50, text)
    c.save()

text_for_pdf = "Новый абзац с измененным шрифтом и размером шрифта."
create_pdf("new_document_formatted.pdf", text_for_pdf)

print("Новый документ с измененным форматированием создан и сохранен в форматах .docx и .pdf.")
# ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ ~ #
'''
1: Импорт библиотек и классов:
'''
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
'''
В этой части кода происходит импорт необходимых библиотек и классов. python-docx используется для работы
с документами Word, а reportlab - для создания PDF.
'''
'''
2: а) Создаем Word файл и заполняем его текстом "Hello Python":
'''
doc = Document()
paragraph = doc.add_paragraph("Hello ")
run = paragraph.add_run("Python")
run.bold = True
doc.save("hello_python.docx")
'''
Создается новый объект Document из python-docx. Добавляется параграф с текстом "Hello " и затем добавляется "запуск"
(run) с текстом "Python", который помечается как жирный. Документ сохраняется в файл "hello_python.docx".
'''
'''
3: б) Открываем существующий документ и выводим жирный текст:
'''
read_doc = Document("hello_python.docx")
bold_text = ""
for paragraph in read_doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text
print("Жирный текст из файла:")
print(bold_text)
'''
Открывается существующий документ "hello_python.docx" и извлекается текст, который был отмечен как жирный.
Этот текст затем выводится на экран.
'''
'''
4: в) Создаем новый документ с абзацем, изменяем шрифт и размер шрифта:
'''
new_doc = Document()
new_paragraph = new_doc.add_paragraph("Новый абзац с измененным шрифтом и размером шрифта.")
new_run = new_paragraph.runs[0]
new_run.font.size = Pt(14)
new_run.font.name = "Arial"
new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
new_doc.save("new_document_formatted.docx")
'''
Создается новый документ, добавляется абзац с текстом, изменяется шрифт и размер шрифта для этого абзаца,
а также устанавливается выравнивание абзаца по центру. Документ сохраняется в файл "new_document_formatted.docx".
'''
'''
5: Создаем PDF документ
'''
def create_pdf(file_name, text):
    pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))  # Замените 'arial.ttf' на путь к файлу вашего шрифта Arial
    c = canvas.Canvas(file_name, pagesize=letter)
    width, height = letter
    c.setFont("Arial", 14)
    c.drawString(width / 2, height - 50, text)
    c.save()

text_for_pdf = "Новый абзац с измененным шрифтом и размером шрифта."
create_pdf("new_document_formatted.pdf", text_for_pdf)
'''
Создается функция create_pdf, которая регистрирует шрифт Arial для reportlab, создает PDF документ с использованием 
canvas, устанавливает шрифт и размер шрифта, добавляет текст и сохраняет документ в файл "new_document_formatted.pdf".
Замените 'arial.ttf' на путь к вашему файлу шрифта Arial.
'''
'''
6: Вывод уведомления на экран:
'''
print("Новый документ с измененным форматированием создан и сохранен в форматах .docx и .pdf.")
'''
Выводится уведомление о том, что новый документ был создан и сохранен в форматах .docx и .pdf.
'''



